from enum import Enum
import os
import random
import json
import argparse

from fastapi import FastAPI, File, UploadFile, Request, Form
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles

from fastapi.templating import Jinja2Templates

from pathlib import Path
from typing import List

from starlette.requests import Request

from pydantic import BaseModel
from typing import List

# import svgwrite
import svgwrite
from svgutils.compose import *
from svgutils.transform import fromstring

from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt

app = FastAPI()


# Initialize the argument parser
parser = argparse.ArgumentParser(description="FastAPI script with command line arguments")

# Define the command line arguments
parser.add_argument("--local",  action="store_true", help="Enter True if running locally, False if running on web")
parser.add_argument("--filename_base", type=str, default="default_filename_base", help="Description for filename_base")
parser.add_argument("--word_file", type=str, default="default_word_file", help="Description for word_file")
parser.add_argument("--rows", type=int, default=20, help="Description for rows")
parser.add_argument("--cols", type=int, default=20, help="Description for cols")

# Parse the command line arguments
args, _ = parser.parse_known_args()


class AlignType(str, Enum):
    left = "LEFT"
    middle = "MIDDLE"
    right = "RIGHT"
class PositionType(str, Enum):
    top = "TOP"
    bottom = "BOTTOM"
class WordInputType(str, Enum):
    file = "FILE"
    list = "LIST"
class TitleParameters(BaseModel):
    align: str | None = AlignType.left
    position: str | None = PositionType.top
    text: str |None = "Word Search"
class WordSearchModel(BaseModel):
    input: str | None = WordInputType.list
    word_list: List[str] | None = None
    # word_file: str | None = None
    row: int | None = 20
    column: int | None = 20
    title: TitleParameters 

# Create a Jinja2Templates instance for rendering HTML templates
templates = Jinja2Templates(directory="templates")

static_folder = "static"

# Mount the static directory
app.mount("/static", StaticFiles(directory="static"), name="static")

# Define the path to the static directory
static_path = Path(__file__).parent / "static"

# Check if the static folder exists, and create it if not
# if not os.path.exists(static_folder):
#     os.makedirs(static_folder)

# Example function that uses the arguments
def example_function(directory_name, arguments):
    # Check if the directory exists
    if not os.path.exists(directory_name):
        # If it doesn't exist, create the directory
        os.makedirs(directory_name)
        print(f"Directory '{directory_name}' created successfully.")
    else:
        print(f"Directory '{directory_name}' already exists.")

    print(f"\nInside example_function:")
    print(f"local: {arguments.local}")
    print(f"filename_base: {arguments.filename_base}")
    print(f"word_file: {arguments.word_file}")
    print(f"rows: {arguments.rows}")
    print(f"cols: {arguments.cols}")

    # Set config variables
    if arguments.filename_base != "default_filename_base":
        docx_output_file = arguments.filename_base +"_doc.docx"
        docx_output_solution_file = arguments.filename_base +"_doc_solution.docx"
        svg_output_file = arguments.filename_base +"_svg.svg"
        svg_output_solution_file = arguments.filename_base +"_svg_solution.svg"
        text_output_file = arguments.filename_base +"_txt.txt"
        text_output_solution_file = arguments.filename_base +"_txt_solution.txt"
    else:
        docx_output_file = "docx-out.docx"
        docx_output_solution_file = "docx-out_solution.docx"
        svg_output_file = "svg_out.svg"
        svg_output_solution_file = "svg_out_solution.svg"
        text_output_file = "text_out.txt"
        text_output_solution_file = "text_out_solution.txt"


    # # Set config variables
    # docx_output_file = "docx-out.docx"
    # docx_output_solution_file = "docx-out_solution.docx"
    # svg_output_file = "svg_out.svg"
    # svg_output_solution_file = "svg_out_solution.svg"
    # text_output_file = "text_out.txt"
    # text_output_solution_file = "text_out_solution.txt"
    font_type = "Courier"
    font_size = "15"
    font_red_value = "0x00"
    font_green_value = "0x00"
    font_blue_value = "0x00"
    display_title = "TRUE"

    title_align = AlignType.left
    title_position = PositionType.top
    # title_text = "Word Search"
    title_text = ""
    word_list = ["CHEESE","BACON","TACOS","SUSHI","PIZZA","DONUTS","BURGER","FRIES","SALAD","GRAPES","YOGURT","PASTA","SANDWICH","SHRIMP","AVOCADO","NACHOS","PANCAKE","CROISSANT","WAFFLE","BAGEL"]
    print("word_list=", word_list)

    # Generate the grids
    # grids = generate_puzzle(data.word_list, data.row, data.column, data.title.text, data.title.align, data.title.position)
    grids = generate_puzzle(word_list, arguments.rows, arguments.cols)
    
    print_docx(word_list, grids[0], grids[1], docx_output_file, docx_output_solution_file,  arguments.rows, arguments.cols, title_text,
           title_align, title_position, font_type, font_size, font_red_value, font_green_value, font_blue_value, display_title)

    # Print the SVG version of the puzzles
    print_svg(word_list, grids[0], grids[1], svg_output_file, svg_output_solution_file, arguments.rows, arguments.cols, title_text, title_align, title_position, display_title)

    # Print the Text version of the puzzles
    print_text(word_list, grids[0], grids[1], text_output_file,
            text_output_solution_file, arguments.rows, arguments.cols, title_text, title_align, title_position, display_title)



@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    # return templates.TemplateResponse("create_wordsearch_grid.html", {"request": request})
    return templates.TemplateResponse("create_wordsearch.html", {"request": request})


@app.get("/svg", response_class=HTMLResponse)
async def render_svg(request: Request):
    return templates.TemplateResponse("svg_template.html", {"request": request})

@app.get("/get_file_links")
async def get_file_links():
    file1_link = "/static/docx-out.docx"
    file2_link = "/static/docx-out_solution.docx"
    file3_link = "/static/svg_out.svg"
    file4_link = "/static/svg_out_solution.svg"
    file5_link = "/static/text_out.txt"
    file6_link = "/static/text_out_solution.txt"

    # Create a dictionary with file paths
    file_paths = {
        "Word Search Puzzle Solution - MS Word version": file1_link,
        "Word Search Puzzle - MS Word version": file2_link,
        "Word Search Puzzle Solution - SVG version": file3_link,
        "Word Search Puzzle - SVG version": file4_link,
        "Word Search Puzzle Solution - Text version": file5_link,
        "Word Search Puzzle - Text version": file6_link
    }

    # Serialize the dictionary to JSON
    response_data = json.dumps(file_paths)

    return response_data

@app.get("/get_files", response_class=HTMLResponse)
async def get_files(request: Request):
    # Define the links to the files
    # file1_link = "/static/docx-out.docx"
    file1_link = f"/static/docx-out.docx"
    print("file1_link=", file1_link)
    file2_link = "/static/docx-out_solution.docx"
    file3_link = "/static/svg_out.svg"
    file4_link = "/static/svg_out_solution.svg"
    file5_link = "/static/text_out.txt"
    file6_link = "/static/text_out_solution.txt"

    # Render the HTML template and pass the file links as context variables
    return templates.TemplateResponse("template.html", {
        "request": request,
        "file1_link": file1_link,
        "file2_link": file2_link,
        "file3_link": file3_link,
        "file4_link": file4_link,
        "file5_link": file5_link,
        "file6_link": file6_link,
    })

@app.get("/get_wordsearch")
async def get_wordsearch():  
    files_to_return = []

    # List of file extensions you want to include
    file_extensions = [".txt", ".svg", ".docx"]

    for root, _, files in os.walk(static_folder):
        for file in files:
            if any(file.endswith(extension) for extension in file_extensions):
                file_path = os.path.join(root, file)
                files_to_return.append(file_path)

    # Return the files as individual responses
    responses = [FileResponse(file_path) for file_path in files_to_return]
    return responses

@app.post("/create_wordsearch")
async def create_wordsearch(data: WordSearchModel, arguments ):    

    # Set config variables
    docx_output_file = "docx-out.docx"
    docx_output_solution_file = "docx-out_solution.docx"
    svg_output_file = "svg_out.svg"
    svg_output_solution_file = "svg_out_solution.svg"
    text_output_file = "text_out.txt"
    text_output_solution_file = "text_out_solution.txt"
    font_type = "Courier"
    font_size = "15"
    font_red_value = "0x00"
    font_green_value = "0x00"
    font_blue_value = "0x00"
    display_title = "TRUE"

    # Generate the grids
    # grids = generate_puzzle(data.word_list, data.row, data.column, data.title.text, data.title.align, data.title.position)
    grids = generate_puzzle(data.word_list, data.row, data.column)
    
    print_docx(data.word_list, grids[0], grids[1], docx_output_file, docx_output_solution_file,  data.row, data.column, data.title.text,
           data.title.align, data.title.position, font_type, font_size, font_red_value, font_green_value, font_blue_value, display_title)

    # Print the SVG version of the puzzles
    print_svg(data.word_list, grids[0], grids[1], svg_output_file, svg_output_solution_file, data.row, data.column, data.title.text, data.title.align, data.title.position, display_title)

    # Print the Text version of the puzzles
    print_text(data.word_list, grids[0], grids[1], text_output_file,
            text_output_solution_file, data.row, data.column, data.title.text, data.title.align, data.title.position, display_title)

    file_links = []

    # Iterate through files in the "static" directory and generate hyperlinks
    # for file_path in static_path.iterdir():
    #     if file_path.is_file():
    #         file_name = file_path.stem  # Use the file name without extension as the hyperlink name
    #         print("file_name=", file_name)
    #         file_links.append({
    #             "filename": file_name,
    #             "url": f"/static/{file_path.name}"
    #         })
    # print("file_links=", file_links)
    return {
        "input": data.input,
        "word_list": data.word_list,
        # "file_name": words_file.filename,
        # "puzzle": grids[0],
        # "puzzle": svg_output_file,
        "puzzle": "This is Text passed from FastAPI on return",
        "row": data.row,
        "column": data.column,
        "title": data.title.model_dump() if data.title else None,
        "svg1": svg_output_file,
        "svg2": svg_output_solution_file,
        "return_svg_file":svg_output_file,
        "return_svg_solution_file":svg_output_solution_file,
        "return_text_file":text_output_file,
        "return_text_solution_file":text_output_solution_file,
        "return_docx_file":docx_output_file,
        "return_docx_solution_file":docx_output_solution_file,
        # "file_links": file_links
    }

def generate_puzzle(words, rows, cols):

    directions = []

    grid = [['-' for _ in range(cols)] for _ in range(rows)]
    placed_grid = [['%' for _ in range(cols)] for _ in range(rows)]

    # Set config variables to DEFAULT add these as ENUMs above
    horz = "TRUE"
    vert = "TRUE"
    diag_right = "TRUE"
    diag_left = "TRUE"
    horz_reverse = "TRUE"
    vert_reverse = "TRUE"

    if horz == "TRUE":
        directions.append('HORIZONTAL')
    if vert == "TRUE":
        directions.append('VERTICAL')
    if (diag_right == "TRUE") or (diag_left == "TRUE"):
        directions.append('DIAGONAL')

    horz_reverse_order = ["FORWARD"]
    if horz_reverse == "TRUE":
        horz_reverse_order.append("REVERSE")

    vert_reverse_order = ["FORWARD"]
    if vert_reverse == "TRUE":
        vert_reverse_order.append("REVERSE")


    for word in words:
        placed = False
        while not placed:
            direction = random.choice(directions)
            order = random.choice(["FORWARD", "REVERSE"])

            if direction == 'HORIZONTAL':
                horz_order = random.choice(horz_reverse_order)
                row = random.randint(0, rows-1)
                col = random.randint(0, cols - len(word))
                if all(grid[row][col + i] in ['-', word[i]] for i in range(len(word))):
                    if horz_order != "FORWARD":
                        # ignore reversing for now
                        # placed_word = reversed(word)
                        placed_word = word
                    else:
                        placed_word = word
                    for i, letter in enumerate(placed_word):
                        grid[row][col + i] = letter.upper()
                        placed_grid[row][col + i] = letter.upper()
                    placed = True

            elif direction == 'VERTICAL':
                vert_order = random.choice(vert_reverse_order)
                row = random.randint(0, rows - len(word))
                col = random.randint(0, cols-1)
                if all(grid[row + i][col] in ['-', word[i]] for i in range(len(word))):
                    if vert_order != "FORWARD":
                        # ignore reversing for now
                        # placed_word = reversed(word)
                        placed_word = word
                    else:
                        placed_word = word
                    for i, letter in enumerate(placed_word):
                        grid[row + i][col] = letter.upper()
                        placed_grid[row + i][col] = letter.upper()
                    placed = True
            else:  # Place Diagional
                row = random.randint(0, rows - len(word))
                col = random.randint(0, cols - len(word))
                drow = random.choice([1, -1])
                dcol = random.choice([1, -1])

                if row + (drow * (len(word) - 1)) < 0 or row + (drow * (len(word) - 1)) >= rows:
                    continue
                if col + (dcol * (len(word) - 1)) < 0 or col + (dcol * (len(word) - 1)) >= cols:
                    continue

                if all(grid[row + drow*i][col + dcol*i] in ['-', word[i]] for i in range(len(word))):
                    if order == "FORWARD":
                        # ignore reversing for now
                        # placed_word = reversed(word)
                        placed_word = word
                    else:
                        placed_word = word
                    for i, letter in enumerate(placed_word):
                        grid[row + drow*i][col + dcol*i] = letter.upper()
                        placed_grid[row + drow*i][col +
                                                dcol*i] = letter.upper()
                    placed = True

    for row in range(rows):
        for col in range(cols):
            if grid[row][col] == '-':
                grid[row][col] = random.choice(['A', 'B', 'C', 'E', 'F', 'G', 'H', 'I', 'J', 'K',
                                                'L', 'M', 'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z'])

    return [grid, placed_grid]

def print_svg(words, grid, placed_grid, svg_filename, svg_solution_filename, rows, cols, title, align, position, display_title):

    width = len(grid[0])
    height = len(grid)
    cell_size = cols

    # display_title = "FALSE"

    # sort the words
    words.sort()

    # Define the number of columns and rows for printing the search words
    num_columns = 3
    col_buffer = (width*width)/3 - cell_size
    start_col_buffer = cell_size * 3
    
    svg_file = os.path.join(static_folder, svg_filename)
    with open(svg_file, "w") as f:
        f.write(f'<svg xmlns="http://www.w3.org/2000/svg" width="{width * cell_size*2}" height="{height * cell_size *2}">\n')
        if (display_title == "TRUE"):
            title_buffer = 1
            title_row = 0
            top_title_buffer = 0
            if (position == PositionType.top):
                top_title_buffer = 1
                match align:
                    case AlignType.left:
                        x = ( cell_size) - (cell_size/2)
                        y = title_row
                        f.write(f'\t<rect x="{x * cell_size}" y="{y * cell_size}" width="{cell_size}" height="{cell_size}" fill-opacity="0" />\n')
                        for i in range(len(title)):
                            f.write(f'\t<text x="{ x + (i * cell_size/2) }" y="{y * cell_size + cell_size / 2}" text-anchor="middle" alignment-baseline="middle" fill="black" font-weight="bold">{title[i]}</text>\n')
                    case AlignType.middle:
                        x = (width/2 - ((len(title)/4)))
                        y=title_row
                        f.write(f'\t<rect x="{x * cell_size}" y="{y * cell_size}" width="{cell_size}" height="{cell_size}" fill-opacity="0" />\n')
                        for i in range(len(title)):
                            f.write(f'\t<text x="{ ((x*cell_size) + (i * cell_size/2))}" y="{y * cell_size + cell_size / 2}" text-anchor="middle" alignment-baseline="middle" fill="black" font-weight="bold">{title[i]}</text>\n')
                    case AlignType.right:
                        x = (width - ((len(title)/2)))
                        y=title_row
                        f.write(f'\t<rect x="{x * cell_size}" y="{y * cell_size}" width="{cell_size}" height="{cell_size}" fill-opacity="0" />\n')
                        for i in range(len(title)):
                            f.write(f'\t<text x="{ ((x*cell_size) + (i * cell_size/2))}" y="{y * cell_size + cell_size / 2}" text-anchor="middle" alignment-baseline="middle" fill="black" font-weight="bold">{title[i]}</text>\n')
                    case _:
                        pass
        else:
            title_buffer = 0

        for y in range(height):
            for x in range(width):
                f.write(
                    f'\t<rect x="{x * cell_size}" y="{(y+top_title_buffer) * cell_size}" width="{cell_size}" height="{cell_size}" fill-opacity="0" />\n')
                f.write(
                    f'\t<text x="{x * cell_size + cell_size / 2}" y="{(y+top_title_buffer) * cell_size + cell_size / 2}" text-anchor="middle" alignment-baseline="middle" fill="black">{grid[y][x]}</text>\n')

        if (position == PositionType.bottom):
            y = y + title_buffer
            match align:
                case AlignType.left:
                    x = ( cell_size) - (cell_size/2)
                    f.write(f'\t<rect x="{x * cell_size}" y="{y * cell_size}" width="{cell_size}" height="{cell_size}" fill-opacity="0" />\n')
                    for i in range(len(title)):
                        f.write(f'\t<text x="{ x + (i * cell_size/2) }" y="{y * cell_size + cell_size / 2}" text-anchor="middle" alignment-baseline="middle" fill="black" font-weight="bold">{title[i]}</text>\n')
                case AlignType.middle:
                    x = (width/2 - ((len(title)/4)))
                    f.write(f'\t<rect x="{x * cell_size}" y="{y * cell_size}" width="{cell_size}" height="{cell_size}" fill-opacity="0" />\n')
                    for i in range(len(title)):
                        f.write(f'\t<text x="{ ((x*cell_size) + (i * cell_size/2))}" y="{y * cell_size + cell_size / 2}" text-anchor="middle" alignment-baseline="middle" fill="black" font-weight="bold">{title[i]}</text>\n')
                case AlignType.right:
                    x = (width - ((len(title)/2)))
                    f.write(f'\t<rect x="{x * cell_size}" y="{y * cell_size}" width="{cell_size}" height="{cell_size}" fill-opacity="0" />\n')
                    for i in range(len(title)):
                        f.write(f'\t<text x="{ ((x*cell_size) + (i * cell_size/2))}" y="{y * cell_size + cell_size / 2}" text-anchor="middle" alignment-baseline="middle" fill="black" font-weight="bold">{title[i]}</text>\n')
                case _:
                    pass

        for i, word in enumerate(words):
            # Calculate the row and column for this word
            row = i // num_columns
            col = i % num_columns
            # set gap for title by adding title_buffer to row if title is inculded
            f.write(f'\t<text x="{(col * (cell_size  + col_buffer)) + start_col_buffer}" y="{(row+ title_buffer) * cell_size + cell_size / 2 + cell_size * cell_size }" text-anchor="middle" alignment-baseline="middle" fill="black" >{word}</text>\n')

        # Close the svg file
        f.write('</svg>\n')
        
    svg_solution_file = os.path.join(static_folder, svg_solution_filename)
    with open(svg_solution_file, "w") as f:
        f.write(f'<svg xmlns="http://www.w3.org/2000/svg" width="{width * cell_size*2}" height="{height * cell_size*2}">\n')

        if (display_title == "TRUE"):
            title_buffer = 1
            title_row = 0
            top_title_buffer = 0
            if (position == PositionType.top):
                top_title_buffer = 1
                match align:
                    case AlignType.left:
                        x = ( cell_size) - (cell_size/2)
                        y = title_row
                        f.write(f'\t<rect x="{x * cell_size}" y="{y * cell_size}" width="{cell_size}" height="{cell_size}" fill-opacity="0" />\n')
                        for i in range(len(title)):
                            f.write(f'\t<text x="{ x + (i * cell_size/2) }" y="{y * cell_size + cell_size / 2}" text-anchor="middle" alignment-baseline="middle" fill="black" font-weight="bold">{title[i]}</text>\n')
                    case AlignType.middle:
                        x = (width/2 - ((len(title)/4)))
                        y=title_row
                        f.write(f'\t<rect x="{x * cell_size}" y="{y * cell_size}" width="{cell_size}" height="{cell_size}" fill-opacity="0" />\n')
                        for i in range(len(title)):
                            f.write(f'\t<text x="{ ((x*cell_size) + (i * cell_size/2))}" y="{y * cell_size + cell_size / 2}" text-anchor="middle" alignment-baseline="middle" fill="black" font-weight="bold">{title[i]}</text>\n')
                    case AlignType.right:
                        x = (width - ((len(title)/2)))
                        y=title_row
                        f.write(f'\t<rect x="{x * cell_size}" y="{y * cell_size}" width="{cell_size}" height="{cell_size}" fill-opacity="0" />\n')
                        for i in range(len(title)):
                            f.write(f'\t<text x="{ ((x*cell_size) + (i * cell_size/2))}" y="{y * cell_size + cell_size / 2}" text-anchor="middle" alignment-baseline="middle" fill="black" font-weight="bold">{title[i]}</text>\n')
                    case _:
                        pass
        else:
            title_buffer = 0

        for y in range(height):
            for x in range(width):
                if (placed_grid[y][x]) == '%':
                    # Not a word so print black
                    f.write(f'\t<rect x="{x * cell_size}" y="{(y+top_title_buffer) * cell_size}" width="{cell_size}" height="{cell_size}" fill-opacity="0" />\n')
                    f.write(f'\t<text x="{x * cell_size + cell_size / 2}" y="{(y+top_title_buffer) * cell_size + cell_size / 2}" text-anchor="middle" alignment-baseline="middle" fill="lightgrey">{grid[y][x]}</text>\n')
                else:
                    # Print solution with formatting
                    f.write(f'\t<rect x="{x * cell_size}" y="{(y+top_title_buffer) * cell_size}" width="{cell_size}" height="{cell_size}" fill-opacity="0" />\n')
                    f.write(f'\t<text x="{x * cell_size + cell_size / 2}" y="{(y+top_title_buffer) * cell_size + cell_size / 2}" text-anchor="middle" alignment-baseline="middle" fill="blue" font-weight = "900">{grid[y][x]}</text>\n')

        if (position == PositionType.bottom):
            y = y + title_buffer
            match align:
                case AlignType.left:
                    x = ( cell_size) - (cell_size/2)
                    f.write(f'\t<rect x="{x * cell_size}" y="{y * cell_size}" width="{cell_size}" height="{cell_size}" fill-opacity="0" />\n')
                    for i in range(len(title)):
                        f.write(f'\t<text x="{ x + (i * cell_size/2) }" y="{y * cell_size + cell_size / 2}" text-anchor="middle" alignment-baseline="middle" fill="black" font-weight="bold">{title[i]}</text>\n')
                case AlignType.middle:
                    x = (width/2 - ((len(title)/4)))
                    f.write(f'\t<rect x="{x * cell_size}" y="{y * cell_size}" width="{cell_size}" height="{cell_size}" fill-opacity="0" />\n')
                    for i in range(len(title)):
                        f.write(f'\t<text x="{ ((x*cell_size) + (i * cell_size/2))}" y="{y * cell_size + cell_size / 2}" text-anchor="middle" alignment-baseline="middle" fill="black" font-weight="bold">{title[i]}</text>\n')
                case AlignType.right:
                    x = (width - ((len(title)/2)))
                    f.write(f'\t<rect x="{x * cell_size}" y="{y * cell_size}" width="{cell_size}" height="{cell_size}" fill-opacity="0" />\n')
                    for i in range(len(title)):
                        f.write(f'\t<text x="{ ((x*cell_size) + (i * cell_size/2))}" y="{y * cell_size + cell_size / 2}" text-anchor="middle" alignment-baseline="middle" fill="black" font-weight="bold">{title[i]}</text>\n')
                case _:
                    pass

        # Loop over the words and add them to the SVG file
        for i, word in enumerate(words):
            # Calculate the row and column for this word
            row = i // num_columns
            col = i % num_columns
            f.write(
                f'\t<text x="{(col * (cell_size  + col_buffer)) + start_col_buffer}" y="{(row+ title_buffer) * cell_size + cell_size / 2 + cell_size * cell_size }" text-anchor="middle" alignment-baseline="middle" fill="blue" font-weight = "900">{word}</text>\n')
        f.write('</svg>\n')
    return [svg_file, svg_solution_file]

def print_docx(words, grid, placed_grid, docx_filename, docx_filename_solution, width, height, title, align, position,  font_type, font_size, red, green, blue, display_title):

    mydoc = Document()
    mydoc_docx = mydoc.add_paragraph()

    style = mydoc.styles['Normal']
    font = style.font
    font.name = font_type
    font.size = Pt(int(font_size))

    words.sort()
    longest_word = ''

    for word in words:
        if len(word) > len(longest_word):
            longest_word = word

    # Check how title is aliigned
    if (align == AlignType.left):
        title_space = 0
    elif (align == AlignType.right):
        title_space = int((int(width)*2 - len(title) - 1))
    else:
        title_space = int((int(width)*2 - len(title)) / 2)

    # Check if Print at TOP or BOTTOM
    if (position == PositionType.top):
        # Print the Title First

        for i in range(title_space):
            doc = mydoc_docx.add_run(" ")

        # print Title
        doc = mydoc_docx.add_run(title)
        doc = mydoc_docx.add_run("\n")

    for x in range(width):
        for y in range(height):
            out = (grid[x][y])
            doc = mydoc_docx.add_run(out)
            doc = mydoc_docx.add_run(" ")
        doc = mydoc_docx.add_run("\n")

    if (position == PositionType.bottom):
        # Print the Title Last

        for i in range(title_space):
            doc = mydoc_docx.add_run(" ")

            # print Title
        doc = mydoc_docx.add_run(title)
        doc = mydoc_docx.add_run("\n")

    # Define the number of columns and rows
    num_columns = 3
    col_buffer = 4

    # Loop over the words and add them to the SVG file
    for i, word in enumerate(words):
        row = i // num_columns
        col = i % num_columns
        doc = mydoc_docx.add_run(word)
        for j in range(col_buffer + (len(longest_word) - len(word))):
            doc = mydoc_docx.add_run(" ")
        if (col == 2):
            doc = mydoc_docx.add_run("\n")

    filename = os.path.join(static_folder, docx_filename)
    mydoc.save(filename)

    # Print the Solution Puzzle

    mydoc_solution = Document()
    mydoc_docx = mydoc_solution.add_paragraph()

    style = mydoc_solution.styles['Normal']
    font = style.font
    font.name = font_type
    font.size = Pt(int(font_size))

    # mydoc.add_page_break()

    # Check if Print at TOP or BOTTOM
    if (position == PositionType.top):
        # Print the Title First

        for i in range(title_space):
            doc = mydoc_docx.add_run(" ")

            # print Title
        doc = mydoc_docx.add_run(title)
        doc = mydoc_docx.add_run("\n")

    for x in range(width):
        for y in range(height):
            if (placed_grid[x][y]) == '%':
                out = (grid[x][y])
                doc = mydoc_docx.add_run(out).bold = False
            else:
                out = (placed_grid[x][y])
                doc = mydoc_docx.add_run(
                    out).font.color.rgb = RGBColor(0, 0, 255)
            doc = mydoc_docx.add_run(" ")
        doc = mydoc_docx.add_run("\n")

    if (position == PositionType.bottom):
        # Print the Title First
        for i in range(title_space):
            doc = mydoc_docx.add_run(" ")

        # print Title
        doc = mydoc_docx.add_run(title)
        doc = mydoc_docx.add_run("\n")

    # Define the number of columns and rows
    num_columns = 3
    col_buffer = 4

    # Loop over the words and add them to the SVG file
    for i, word in enumerate(words):
        row = i // num_columns
        col = i % num_columns
        doc = mydoc_docx.add_run(word)
        for j in range(col_buffer + (len(longest_word) - len(word))):
            doc = mydoc_docx.add_run(" ")
        if (col == 2):
            doc = mydoc_docx.add_run("\n")

    # get the filename only
    filename = os.path.join(static_folder, docx_filename_solution)
    mydoc_solution.save(filename)

def print_text(words, grid, placed_grid, text_filename, test_solution_filename, rows, cols, title, align, position, display_title):

    words.sort()
    longest_word = ''
    for word in words:
        if len(word) > len(longest_word):
            longest_word = word

    # create text filie
    text_file = os.path.join(static_folder, text_filename)

    # create text file for solution
    text_solution_file = os.path.join(static_folder, test_solution_filename)

    # Define the number of columns and rows
    num_columns = 3
    col_buffer = 4

    # Check how title is aliigned
    if (align == AlignType.left):
        title_space = 0
    elif (align == AlignType.right):
        title_space = int((int(cols)*2 - len(title) - 1))
    else:
        title_space = int((int(cols)*2 - len(title)) / 2)


    with open(text_file, 'w') as f:

        # Check if Print at TOP 
        if (position == PositionType.top):
            # Print the Title at top
            for y in range(title_space):
                f.write(" ")
            f.write(title)
            f.write('\n')

        for x in range(rows):
            for y in range(cols):
                f.write(grid[x][y])
                f.write(" ")
            f.write('\n')

        # Check if Print at TOP or BOTTOM
        if (position == PositionType.bottom):
            # Print the Title at bottom 
            for y in range(title_space):
                f.write(" ")
            f.write(title)
            f.write('\n')
            f.write('\n')

        # Loop over the words and add them to the text file
        for i, word in enumerate(words):
            row = i // num_columns
            col = i % num_columns
            f.write(word)
            for j in range(col_buffer + (len(longest_word) - len(word))):
                f.write(" ")
            if (col == 2):
                f.write("\n")

    with open(text_solution_file, 'w') as f:

        for x in range(rows):
            for y in range(cols):
                if (placed_grid[x][y] == "%"):
                    f.write(" ")
                else:
                    f.write(placed_grid[x][y])
                f.write(" ")

            f.write('\n')

        # Loop over the words and add them to the SVG file
        for i, word in enumerate(words):
            row = i // num_columns
            col = i % num_columns
            f.write(word)
            for j in range(col_buffer + (len(longest_word) - len(word))):
                f.write(" ")
            if (col == 2):
                f.write("\n")


# Run the FastAPI app
if __name__ == "__main__":
    import uvicorn
    print("args=", args)
    if args.local:
        # Specify the directory name you want to create
        directory_name = "new_directory"    

        example_function(directory_name, args)    

        # Call the function to create the directory
        # create_directory(directory_name)
    else:
        uvicorn.run(app, host="127.0.0.1", port=8000)
